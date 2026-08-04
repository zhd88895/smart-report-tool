# -*- coding: utf-8 -*-
"""
将 Markdown 用户使用说明文档转换为 Word (.docx) 和 HTML 两种格式。
"""

import re
import os
from pathlib import Path

# ============================================================
# 配置
# ============================================================
BASE_DIR = Path(__file__).parent
MD_FILE = BASE_DIR / "用户使用说明文档.md"
DOCX_FILE = BASE_DIR / "用户使用说明文档.docx"
HTML_FILE = BASE_DIR / "用户使用说明文档.html"

# ============================================================
# Part 1: 生成 Word 文档
# ============================================================
def generate_docx():
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # ── 全局样式设置 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    # 设置中文字体
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题样式
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        hfont = heading_style.font
        hfont.name = '微软雅黑'
        hfont.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
        hfont.bold = True
        hPr = heading_style.element.get_or_add_rPr()
        hFonts = hPr.find(qn('w:rFonts'))
        if hFonts is None:
            hFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
            hPr.insert(0, hFonts)
        else:
            hFonts.set(qn('w:eastAsia'), '微软雅黑')

        if level == 1:
            hfont.size = Pt(18)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            hfont.size = Pt(14)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        else:
            hfont.size = Pt(12)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── 读取 Markdown 并解析 ──
    md_text = MD_FILE.read_text(encoding='utf-8')
    lines = md_text.split('\n')

    # ── 封面页 ──
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('智能报告生成工具')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run('用户使用说明文档')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    for _ in range(3):
        doc.add_paragraph()

    info_items = [
        ('文档版本', 'v1.0'),
        ('适用系统版本', 'v0.3.2'),
        ('编写日期', '2026-07-21'),
        ('文档类型', '用户使用说明'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_page_break()

    # ── 逐行解析 Markdown 并写入 ──
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    in_blockquote = False

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        # 解析表格行
        data_rows = []
        for row_text in table_rows:
            cells = [c.strip() for c in row_text.strip('|').split('|')]
            # 跳过分隔行（如 |---|---|）
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            data_rows.append(cells)

        if not data_rows:
            table_rows = []
            in_table = False
            return

        num_cols = max(len(r) for r in data_rows)
        table = doc.add_table(rows=len(data_rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row_data in enumerate(data_rows):
            for ci, cell_text in enumerate(row_data):
                if ci < num_cols:
                    cell = table.rows[ri].cells[ci]
                    # 清除 Markdown 加粗标记
                    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                    clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
                    cell.text = clean_text
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(2)
                        paragraph.paragraph_format.space_before = Pt(2)
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                            run.font.name = '微软雅黑'
                            rpr = run._element.get_or_add_rPr()
                            rf = rpr.find(qn('w:rFonts'))
                            if rf is None:
                                rf = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微软雅黑"/>')
                                rpr.insert(0, rf)
                            else:
                                rf.set(qn('w:eastAsia'), '微软雅黑')

            # 表头行加粗 + 背景色
            if ri == 0:
                for ci in range(num_cols):
                    for paragraph in table.rows[ri].cells[ci].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                    # 浅灰背景
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F5E9" w:val="clear"/>')
                    table.rows[ri].cells[ci]._element.get_or_add_tcPr().append(shading)

        table_rows = []
        in_table = False

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        # 灰色背景效果通过缩进+字体模拟
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        # 添加底纹
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        run._element.get_or_add_rPr().append(shading)
        code_lines = []

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 表格
        if '|' in line and line.strip().startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()

        # 空行
        if not line.strip():
            i += 1
            continue

        # 分隔线
        if line.strip() == '---':
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # 清除 Markdown 标记
            clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
            doc.add_heading(clean_text, level=level)
            i += 1
            continue

        # 引用块
        if line.strip().startswith('>'):
            quote_text = line.strip().lstrip('> ').strip()
            # 清除 Markdown 标记
            clean_quote = re.sub(r'\*\*(.+?)\*\*', r'\1', quote_text)
            clean_quote = re.sub(r'`(.+?)`', r'\1', clean_quote)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f'💡 {clean_quote}')
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            i += 1
            continue

        # 无序列表
        if re.match(r'^[-*]\s+', line.strip()):
            text = re.sub(r'^[-*]\s+', '', line.strip())
            clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
            p = doc.add_paragraph(style='List Bullet')
            # 处理内联格式
            parts = re.split(r'(\*\*.+?\*\*|`.+?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    p.add_run(part)
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.+?\*\*|`.+?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    p.add_run(part)
            i += 1
            continue

        # 普通段落
        clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        clean_line = re.sub(r'`(.+?)`', r'\1', clean_line)
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.+?\*\*|`.+?`)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            else:
                p.add_run(part)
        i += 1

    # 最后刷新
    if in_table:
        flush_table()
    if in_code_block:
        flush_code()

    # ── 免责声明 ──
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('免责声明')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('本文档由 AI 辅助生成，重要决策请经专业人员核验。\n文档内容基于项目源码 v0.3.2 版本分析，如有更新请以实际代码为准。')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.save(str(DOCX_FILE))
    print(f"[OK] Word 文档已生成: {DOCX_FILE}")


# ============================================================
# Part 2: 生成 HTML 网页
# ============================================================
def generate_html():
    import markdown

    md_text = MD_FILE.read_text(encoding='utf-8')

    # 使用 markdown 扩展
    html_body = markdown.markdown(
        md_text,
        extensions=[
            'tables',
            'fenced_code',
            'codehilite',
            'toc',
            'nl2br',
        ],
        extension_configs={
            'codehilite': {'css_class': 'highlight', 'guess_lang': False},
        }
    )

    html_template = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>智能报告生成工具 — 用户使用说明文档</title>
    <style>
        :root {{
            --primary: #0F766E;
            --primary-light: #14B8A6;
            --primary-bg: #F0FDFA;
            --text: #1E293B;
            --text-secondary: #64748B;
            --bg: #F8FAFC;
            --card: #FFFFFF;
            --border: #E2E8F0;
            --code-bg: #F1F5F9;
            --table-header: #E8F5E9;
        }}

        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', '微软雅黑', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.8;
            color: var(--text);
            background: var(--bg);
        }}

        /* ── 侧边导航 ── */
        .sidebar {{
            position: fixed;
            top: 0;
            left: 0;
            width: 280px;
            height: 100vh;
            overflow-y: auto;
            background: var(--card);
            border-right: 1px solid var(--border);
            padding: 24px 16px;
            z-index: 100;
        }}

        .sidebar-title {{
            font-size: 16px;
            font-weight: 700;
            color: var(--primary);
            margin-bottom: 20px;
            padding-bottom: 12px;
            border-bottom: 2px solid var(--primary-light);
        }}

        .sidebar ul {{
            list-style: none;
            padding: 0;
        }}

        .sidebar li {{
            margin-bottom: 4px;
        }}

        .sidebar a {{
            display: block;
            padding: 6px 12px;
            color: var(--text-secondary);
            text-decoration: none;
            font-size: 14px;
            border-radius: 6px;
            transition: all 0.2s;
        }}

        .sidebar a:hover {{
            background: var(--primary-bg);
            color: var(--primary);
        }}

        .sidebar .toc-h3 {{
            padding-left: 24px;
            font-size: 13px;
        }}

        /* ── 主内容区 ── */
        .main-content {{
            margin-left: 280px;
            max-width: 960px;
            padding: 40px 48px 80px;
        }}

        /* ── 封面 ── */
        .cover {{
            text-align: center;
            padding: 80px 0 60px;
            margin-bottom: 40px;
            border-bottom: 3px solid var(--primary);
        }}

        .cover h1 {{
            font-size: 36px;
            color: var(--primary);
            margin-bottom: 8px;
        }}

        .cover .subtitle {{
            font-size: 22px;
            color: var(--primary-light);
            margin-bottom: 40px;
        }}

        .cover .meta {{
            color: var(--text-secondary);
            font-size: 15px;
        }}

        .cover .meta span {{
            margin: 0 16px;
        }}

        /* ── 标题 ── */
        h1 {{
            font-size: 28px;
            color: var(--primary);
            margin: 48px 0 20px;
            padding-bottom: 10px;
            border-bottom: 2px solid var(--primary-light);
        }}

        h2 {{
            font-size: 22px;
            color: var(--primary);
            margin: 36px 0 16px;
            padding-left: 12px;
            border-left: 4px solid var(--primary-light);
        }}

        h3 {{
            font-size: 17px;
            color: var(--text);
            margin: 24px 0 12px;
        }}

        h4 {{
            font-size: 15px;
            color: var(--text);
            margin: 20px 0 10px;
        }}

        /* ── 段落 ── */
        p {{
            margin-bottom: 14px;
            text-align: justify;
        }}

        /* ── 链接 ── */
        a {{
            color: var(--primary);
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        /* ── 加粗 ── */
        strong {{
            color: var(--text);
        }}

        /* ── 列表 ── */
        ul, ol {{
            margin: 8px 0 16px 24px;
        }}

        li {{
            margin-bottom: 6px;
        }}

        /* ── 代码块 ── */
        pre {{
            background: var(--code-bg);
            border: 1px solid var(--border);
            border-radius: 8px;
            padding: 16px 20px;
            overflow-x: auto;
            margin: 12px 0 20px;
            font-size: 13px;
            line-height: 1.6;
        }}

        pre code {{
            background: none;
            padding: 0;
            border: none;
            font-family: 'Cascadia Code', 'Fira Code', 'Consolas', 'Monaco', monospace;
        }}

        /* ── 内联代码 ── */
        code {{
            background: var(--code-bg);
            padding: 2px 6px;
            border-radius: 4px;
            font-family: 'Cascadia Code', 'Fira Code', 'Consolas', 'Monaco', monospace;
            font-size: 0.9em;
            color: #D63384;
        }}

        /* ── 表格 ── */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 16px 0 24px;
            font-size: 14px;
        }}

        th {{
            background: var(--table-header);
            color: var(--primary);
            font-weight: 600;
            text-align: left;
            padding: 10px 14px;
            border: 1px solid var(--border);
        }}

        td {{
            padding: 9px 14px;
            border: 1px solid var(--border);
            vertical-align: top;
        }}

        tr:hover td {{
            background: var(--primary-bg);
        }}

        /* ── 引用块 ── */
        blockquote {{
            border-left: 4px solid var(--primary-light);
            background: var(--primary-bg);
            padding: 12px 20px;
            margin: 16px 0;
            border-radius: 0 8px 8px 0;
            color: var(--text-secondary);
            font-style: italic;
        }}

        /* ── 分隔线 ── */
        hr {{
            border: none;
            height: 1px;
            background: var(--border);
            margin: 32px 0;
        }}

        /* ── 回到顶部按钮 ── */
        .back-to-top {{
            position: fixed;
            bottom: 32px;
            right: 32px;
            width: 44px;
            height: 44px;
            background: var(--primary);
            color: white;
            border: none;
            border-radius: 50%;
            cursor: pointer;
            font-size: 20px;
            display: none;
            align-items: center;
            justify-content: center;
            box-shadow: 0 4px 12px rgba(15, 118, 110, 0.3);
            transition: transform 0.2s;
            z-index: 200;
        }}

        .back-to-top:hover {{
            transform: scale(1.1);
        }}

        /* ── 打印样式 ── */
        @media print {{
            .sidebar, .back-to-top {{
                display: none !important;
            }}
            .main-content {{
                margin-left: 0;
            }}
            body {{
                background: white;
            }}
        }}

        /* ── 响应式 ── */
        @media (max-width: 1024px) {{
            .sidebar {{
                display: none;
            }}
            .main-content {{
                margin-left: 0;
                padding: 24px 20px;
            }}
        }}
    </style>
</head>
<body>
    <!-- 侧边导航 -->
    <nav class="sidebar" id="sidebar">
        <div class="sidebar-title">📋 目录导航</div>
        <ul id="toc-list">
        </ul>
    </nav>

    <!-- 主内容 -->
    <div class="main-content">
        <div class="cover">
            <h1>智能报告生成工具</h1>
            <div class="subtitle">用户使用说明文档</div>
            <div class="meta">
                <span>文档版本：v1.0</span>
                <span>适用系统：v0.3.2</span>
                <span>日期：2026-07-21</span>
            </div>
        </div>

        {html_body}

        <hr>
        <blockquote>
            <strong>免责声明</strong>：本文档由 AI 辅助生成，重要决策请经专业人员核验。文档内容基于项目源码 v0.3.2 版本分析，如有更新请以实际代码为准。
        </blockquote>
    </div>

    <!-- 回到顶部 -->
    <button class="back-to-top" id="backToTop" onclick="window.scrollTo({{top:0,behavior:'smooth'}})">↑</button>

    <script>
        // 自动生成侧边目录
        (function() {{
            const tocList = document.getElementById('toc-list');
            const headings = document.querySelectorAll('.main-content h1, .main-content h2, .main-content h3');
            const fragment = document.createDocumentFragment();

            headings.forEach((h, index) => {{
                const id = 'heading-' + index;
                h.id = id;

                const li = document.createElement('li');
                const a = document.createElement('a');
                a.href = '#' + id;
                a.textContent = h.textContent;

                if (h.tagName === 'H3') {{
                    li.className = 'toc-h3';
                }}

                li.appendChild(a);
                fragment.appendChild(li);
            }});

            tocList.appendChild(fragment);
        }})();

        // 回到顶部按钮
        window.addEventListener('scroll', function() {{
            const btn = document.getElementById('backToTop');
            btn.style.display = window.scrollY > 300 ? 'flex' : 'none';
        }});

        // 平滑滚动
        document.querySelectorAll('.sidebar a').forEach(a => {{
            a.addEventListener('click', function(e) {{
                e.preventDefault();
                const target = document.querySelector(this.getAttribute('href'));
                if (target) {{
                    target.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
                }}
            }});
        }});
    </script>
</body>
</html>'''

    HTML_FILE.write_text(html_template, encoding='utf-8')
    print(f"[OK] HTML 网页已生成: {HTML_FILE}")


# ============================================================
# 主程序
# ============================================================
if __name__ == '__main__':
    print("=" * 50)
    print("  文档格式转换工具")
    print("=" * 50)

    print("\n[1/2] 生成 Word 文档...")
    try:
        generate_docx()
    except Exception as e:
        print(f"[ERROR] Word 生成失败: {e}")

    print("\n[2/2] 生成 HTML 网页...")
    try:
        generate_html()
    except Exception as e:
        print(f"[ERROR] HTML 生成失败: {e}")

    print("\n" + "=" * 50)
    print("  转换完成！")
    print("=" * 50)
